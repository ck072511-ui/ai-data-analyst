import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from typing import Any, Dict
from docx import Document
from docx.shared import Inches

class DOCXReportGenerator:
    @staticmethod
    def generate(file_path: str, data: Dict[str, Any], branding: Dict[str, Any]):
        """Generates an editable DOCX report using python-docx structures."""
        doc = Document()

        company_name = branding.get("company_name", "Enterprise Systems")
        version = branding.get("version", "1.0.0")

        # Title Section
        doc.add_heading(data.get("title", "Enterprise Analytics Report"), level=0)
        doc.add_paragraph(f"Compiled for: {company_name} | Version: {version}")
        doc.add_paragraph(f"Timestamp: {data.get('timestamp')}")
        doc.add_paragraph().paragraph_format.space_after = Inches(0.5)

        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(data.get("final_answer", "Analyzed dataset metrics successfully."))

        # Key Performance Indicators
        doc.add_heading("Key Performance Indicators (KPIs)", level=1)
        kpis = data.get("kpis", [])
        if kpis:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Shading Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Indicator'
            hdr_cells[1].text = 'Value'
            for k in kpis[:4]:
                row_cells = table.add_row().cells
                row_cells[0].text = k.get("title", "")
                row_cells[1].text = str(k.get("value", ""))
        else:
            doc.add_paragraph("No KPIs recommended for this dataset view.")

        # Embedded Chart Image
        doc.add_heading("Visualized Trend Graphic", level=1)
        chart_path = f"{file_path}.png"
        try:
            plt.figure(figsize=(6, 3))
            plt.plot([1, 2, 3, 4], [10, 20, 25, 30], label="Sales Value", marker='s', color='#2563eb')
            plt.title("Sample Performance Metrics Trend")
            plt.grid(True)
            plt.legend()
            plt.tight_layout()
            plt.savefig(chart_path)
            plt.close()
            
            doc.add_picture(chart_path, width=Inches(5))
        except Exception as e:
            doc.add_paragraph(f"Failed to generate and embed Matplotlib graphic: {e}")

        # AI Insights
        doc.add_heading("AI Business Insights & SQL Context", level=1)
        insights = data.get("insights", [])
        for ins in insights[:3]:
            doc.add_paragraph(ins, style='List Bullet')
        
        doc.add_heading("Executed SQL Query", level=2)
        doc.add_paragraph(data.get("sql", "No SQL executed."))

        # Citations
        doc.add_heading("Cited Source Documents", level=1)
        citations = data.get("citations", [])
        if citations:
            for c in citations[:2]:
                doc.add_paragraph(f"{c.get('filename')} (Pg {c.get('page_number')}): {c.get('text_content')}")
        else:
            doc.add_paragraph("No offline document citations cited.")

        # Save document
        doc.save(file_path)

        # Cleanup chart image
        if os.path.exists(chart_path):
            try:
                os.remove(chart_path)
            except:
                pass
