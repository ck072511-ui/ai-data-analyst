import os
import logging
import pandas as pd
from typing import List, Tuple

logger = logging.getLogger(__name__)

class DocumentIngestionService:
    @staticmethod
    def validate_file(file_path: str, extension: str) -> bool:
        """Verifies if the file exists and is not corrupted by trying to parse basic structure."""
        if not os.path.exists(file_path):
            return False
        
        try:
            if extension == "pdf":
                import pypdf
                reader = pypdf.PdfReader(file_path)
                # Access pages count to verify integrity
                _ = len(reader.pages)
            elif extension == "docx":
                import docx
                _ = docx.Document(file_path)
            elif extension == "csv":
                _ = pd.read_csv(file_path, nrows=2)
            elif extension in ["txt", "md", "markdown"]:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    _ = f.read(100)
            return True
        except Exception as e:
            logger.warning(f"File validation failure for {file_path}: {e}")
            return False

    @staticmethod
    def extract_text(file_path: str, extension: str) -> List[Tuple[int, str]]:
        """Extracts text pages/sections mapping them to page numbers (1-indexed)."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        pages = []
        try:
            if extension == "pdf":
                try:
                    import pypdf
                    reader = pypdf.PdfReader(file_path)
                    for idx, page in enumerate(reader.pages):
                        text = page.extract_text() or ""
                        pages.append((idx + 1, text))
                except ImportError:
                    logger.error("pypdf is missing. Cannot parse PDF.")
                    raise ImportError("Missing pypdf dependency for PDF ingestion.")

            elif extension == "docx":
                try:
                    import docx
                    doc = docx.Document(file_path)
                    full_text = []
                    for para in doc.paragraphs:
                        if para.text:
                            full_text.append(para.text)
                    pages.append((1, "\n".join(full_text)))
                except ImportError:
                    logger.error("docx dependency is missing.")
                    raise ImportError("Missing python-docx dependency for DOCX ingestion.")

            elif extension == "csv":
                df = pd.read_csv(file_path)
                pages.append((1, df.to_string(index=False)))

            elif extension in ["txt", "md", "markdown"]:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
                pages.append((1, content))
            else:
                raise ValueError(f"Ingestion format not supported: {extension}")

        except Exception as e:
            logger.exception(f"Text extraction failed for {file_path}")
            raise e

        return pages
